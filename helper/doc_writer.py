import os
from PIL import Image
from typing import Union, BinaryIO, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.image.exceptions import UnrecognizedImageError



class DocumentWriter:

    def __init__(self,output):
        self.__doc = Document()
        self.document_name = output

    # Writing the header of the document
    def write_header(self,heading: str):
        head = self.__doc.add_heading()
        head_run = head.add_run(heading)
        head_run.font.size = Pt(18)
        head_run.font.color.rgb = RGBColor(0,0,0)

    # Writing a paragraph
    def write_paragraph(self,paragraph: str):
        para = self.__doc.add_paragraph()
        para_run = para.add_run(paragraph)
        para_run.font.size = Pt(13)
        para_run.font.name = 'Calibri'
        para_format = para.paragraph_format
        #para_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Adding a picture to the document
    def add_image(self,
                  file: Union[str, BinaryIO],
                  width: Union[int, float] = 6,
                  height: Union[int, float] = 4.1):
        try:
            self.__doc.add_picture(file,
                                   width=Inches(width),
                                   height=Inches(height)
                )
        # Handling UnrecognizedImageError
        except UnrecognizedImageError:
            Image.open(file).save(file)
            self.__doc.add_picture(file,
                                   width=Inches(width),
                                   height=Inches(height)
                )
    # Save the Document as .docx file
    def save_document(self):
        self.__doc.save(self.document_name)

    

# For Write Documents that are realted to nasaspaceflight.com
class NasaSpaceFlightDocumentWriter(DocumentWriter):
    
    def __init__(self, document_name):
        super().__init__(document_name)

    def write_document(
            self,
            header: str,
            lead_image: Union[str, os.PathLike],
            paragraph: str,
            image_count: int,
            image_prefix: str
        ):
        # Writing the heading of the Article
        self.write_header(header)
        # Adding the lead image 
        self.add_image(lead_image)
        # Writing the paragraph
        self.write_paragraph(paragraph)

        # Adding remaining images at the end of the document

        for pic_num in range(2,image_count):
            img = f'src/{image_prefix}-pic-{pic_num}.jpg'
            self.add_image(img)
        
        # Finally saving the document(.docx) at the given path
        self.save_document()
        
